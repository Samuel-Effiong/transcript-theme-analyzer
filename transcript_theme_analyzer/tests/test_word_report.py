import os

from docx import Document

from transcript_theme_analyzer.report import build_report
from transcript_theme_analyzer.word_report import build_docx, write_docx, write_docx_from_data


def make_success(score, reasoning="r", locations=None, model_used="m1", chunked=False):
    return {
        "theme": "t",
        "relevance_score": score,
        "reasoning": reasoning,
        "locations": locations or [],
        "model_used": model_used,
        "chunked": chunked,
        "_meta": {"elapsed_seconds": 1.0, "usage": {"total_tokens": 100}},
    }


def make_error(model, error="boom"):
    return {"model": model, "error": error}


def _headings(document, level):
    return [p.text for p in document.paragraphs if p.style.name == f"Heading {level}"]


def test_document_title_is_the_theme():
    transcripts = {"ep1": [("m1", make_success(80, locations=[
        {"excerpt": "e1", "title": "First topic"},
    ]))]}
    report = build_report("the glory of God", transcripts)
    document = build_docx(report)
    assert _headings(document, 1) == ["the glory of God"]


def test_one_section_per_transcript_named_after_transcript():
    transcripts = {
        "Sunday Service": [("m1", make_success(80, locations=[
            {"excerpt": "e1", "title": "Topic A"},
        ]))],
        "Leadership Conference": [("m1", make_success(70, locations=[
            {"excerpt": "e2", "title": "Topic B"},
        ]))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)
    headings2 = _headings(document, 2)
    assert "Transcript: Sunday Service" in headings2
    assert "Transcript: Leadership Conference" in headings2


def test_excerpt_heading_uses_context_not_the_raw_keyword():
    transcripts = {
        "ep1": [("m1", make_success(80, locations=[
            {"excerpt": "e1", "title": "God's Glory Revealed Through Obedience."},
            {"excerpt": "e2", "title": "The Purpose of Living for God's Glory."},
        ]))],
    }
    report = build_report("the glory of God", transcripts)
    document = build_docx(report)
    headings3 = _headings(document, 3)
    assert "God's Glory Revealed Through Obedience" in headings3
    assert "The Purpose of Living for God's Glory" in headings3
    # Distinct excerpts get distinct headings, not a repeat of the theme.
    assert len(set(headings3)) == 2


def test_multiple_excerpts_under_same_transcript_all_included():
    transcripts = {
        "ep1": [("m1", make_success(80, locations=[
            {"excerpt": "e1", "title": "Topic A"},
            {"excerpt": "e2", "title": "Topic B"},
            {"excerpt": "e3", "title": "Topic C"},
        ]))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)
    body_text = "\n".join(p.text for p in document.paragraphs)
    assert "e1" in body_text and "e2" in body_text and "e3" in body_text
    assert len(_headings(document, 3)) == 3


def test_transcripts_with_no_located_excerpts_are_skipped():
    transcripts = {
        "has_locations": [("m1", make_success(80, locations=[
            {"excerpt": "e1", "title": "Topic A"},
        ]))],
        "no_locations": [("m1", make_success(50, locations=[]))],
        "all_failed": [("m1", make_error("m1"))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)
    headings2 = _headings(document, 2)
    assert headings2 == ["Transcript: has_locations"]


def test_min_score_filters_out_low_scoring_transcripts():
    transcripts = {
        "high": [("m1", make_success(90, locations=[
            {"excerpt": "e1", "title": "Topic A"},
        ]))],
        "low": [("m1", make_success(10, locations=[
            {"excerpt": "e2", "title": "Topic B"},
        ]))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report, min_score=50)
    assert _headings(document, 2) == ["Transcript: high"]


def test_no_matching_transcripts_produces_a_document_without_crashing():
    transcripts = {"ep1": [("m1", make_success(80, locations=[]))]}
    report = build_report("theme", transcripts)
    document = build_docx(report)
    assert _headings(document, 2) == []
    assert any("No transcripts" in p.text for p in document.paragraphs)


def test_multi_paragraph_excerpt_gets_open_quote_only_on_first_paragraph_and_close_on_last():
    excerpt = "First paragraph of the passage.\n\nSecond paragraph.\n\nThird and final paragraph."
    transcripts = {
        "ep1": [("m1", make_success(80, locations=[
            {"excerpt": excerpt, "title": "A Long Passage"},
        ]))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)
    body_paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    assert "“First paragraph of the passage." in body_paragraphs
    assert "Second paragraph." in body_paragraphs
    assert "Third and final paragraph.”" in body_paragraphs
    # The middle paragraph carries neither a leading nor trailing quote mark.
    assert "“Second paragraph.”" not in body_paragraphs
    assert "“Second paragraph." not in body_paragraphs


def test_write_docx_produces_a_valid_openable_file(tmp_path):
    transcripts = {"ep1": [("m1", make_success(80, locations=[
        {"excerpt": "e1", "title": "Topic A", "timestamp": "[01:00]", "speaker": "Host"},
    ]))]}
    out_path = str(tmp_path / "out.docx")
    write_docx_from_data("theme", transcripts, out_path)
    assert os.path.isfile(out_path)
    reopened = Document(out_path)
    assert reopened.paragraphs[0].text == "theme"


def test_hostile_content_is_stored_literally_not_executed(tmp_path):
    hostile = "=cmd|' /C calc'!A1"  # CSV-injection-shaped string; docx just stores text
    transcripts = {
        "ep1": [("m1", make_success(80, locations=[
            {"excerpt": hostile, "title": hostile},
        ]))],
    }
    out_path = str(tmp_path / "out.docx")
    write_docx_from_data("theme", transcripts, out_path)
    reopened = Document(out_path)
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    assert hostile in all_text


def test_overview_table_lists_every_transcript_with_its_score_in_ranked_order():
    transcripts = {
        "high": [("m1", make_success(90, locations=[
            {"excerpt": "e1", "title": "Topic A"},
        ]))],
        "no_excerpts": [("m1", make_success(50, locations=[]))],
        "failed": [("m1", make_error("m1"))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)
    assert len(document.tables) == 1
    table = document.tables[0]

    # header + one row per transcript
    assert len(table.rows) == 1 + len(transcripts)
    header_texts = [c.text for c in table.rows[0].cells]
    assert header_texts == ["Transcript", "Relevance Score"]

    body_rows = [(row.cells[0].text, row.cells[1].text) for row in table.rows[1:]]
    assert body_rows[0] == ("high", "90")
    assert body_rows[1][0].startswith("no_excerpts")
    assert body_rows[1][1] == "50"
    assert body_rows[2][0].startswith("failed")
    assert body_rows[2][1] == "FAILED"


def test_overview_table_links_only_transcripts_with_a_detail_section():
    transcripts = {
        "has_section": [("m1", make_success(80, locations=[
            {"excerpt": "e1", "title": "Topic A"},
        ]))],
        "no_section": [("m1", make_success(20, locations=[]))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)
    table = document.tables[0]

    linked_row_xml = table.rows[1].cells[0].paragraphs[0]._p.xml
    unlinked_row_xml = table.rows[2].cells[0].paragraphs[0]._p.xml

    assert "w:hyperlink" in linked_row_xml
    assert 'w:anchor="t0"' in linked_row_xml
    assert "w:hyperlink" not in unlinked_row_xml


def test_transcript_heading_carries_the_matching_bookmark():
    transcripts = {
        "only_one": [("m1", make_success(80, locations=[
            {"excerpt": "e1", "title": "Topic A"},
        ]))],
    }
    report = build_report("theme", transcripts)
    document = build_docx(report)

    heading_paragraph = next(
        p for p in document.paragraphs if p.style.name == "Heading 2"
    )
    assert 'w:name="t0"' in heading_paragraph._p.xml
