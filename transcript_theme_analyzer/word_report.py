"""Generates a single consolidated Word (.docx) document with every relevant
excerpt from every transcript that matched the search theme:

- Document title = the search theme.
- An overview table (transcript × relevance score, in ranked order) right
  after the title, with clickable links that jump to each transcript's
  section further down.
- One section per transcript with at least one located excerpt, headed
  ``Transcript: <name>`` so an excerpt can always be traced back to its
  source. One sub-heading per excerpt, auto-generated from that excerpt's
  own context rather than just the search keyword.

Usage (regenerate later from the internal cache, without re-running analysis):
    python -m transcript_theme_analyzer.word_report --cache results/.raw_results.json
"""
from __future__ import annotations

import argparse
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .cache import read_cache
from .report import BatchReport, TranscriptResult, build_report

_LINK_COLOR = RGBColor(0x25, 0x6A, 0xBF)  # same blue used for html links/bars
_FAILED_COLOR = RGBColor(0xD0, 0x3B, 0x3B)  # same red used for the html FAILED badge


def _add_horizontal_rule(document: Document) -> None:
    """A bottom-bordered empty paragraph, python-docx's standard recipe for
    a horizontal rule (Word has no native <hr> element)."""
    p = document.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text: str, bookmark_name: str) -> None:
    """Appends a clickable link to ``paragraph`` that jumps to ``bookmark_name``
    elsewhere in the same document (an anchor-only hyperlink — no external
    relationship, so python-docx's public hyperlink-reading API won't see it,
    but Word/LibreOffice render and follow it like any other link)."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "256ABF")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run_element.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)

    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _add_excerpt(document: Document, location: dict) -> None:
    # The heading is generated from the location's own title (produced by
    # the analysis pass for that specific passage), not from the search
    # theme, so headings differ across excerpts even within the same
    # transcript.
    heading_text = (location.get("title") or "").strip().rstrip(".")
    document.add_heading(heading_text or "Relevant Passage", level=3)

    meta_bits = [b for b in (location.get("timestamp"), location.get("speaker")) if b]
    if meta_bits:
        meta_p = document.add_paragraph(" · ".join(meta_bits))
        for run in meta_p.runs:
            run.italic = True
            run.font.size = Pt(9)

    excerpt = (location.get("excerpt") or "").strip()
    paragraphs = [p.strip() for p in excerpt.split("\n\n") if p.strip()] or [excerpt]
    for i, para_text in enumerate(paragraphs):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        prefix = "“" if i == 0 else ""
        suffix = "”" if i == len(paragraphs) - 1 else ""
        p.add_run(f"{prefix}{para_text}{suffix}").italic = True


def _matching_transcripts(report: BatchReport, min_score: int | None) -> list[TranscriptResult]:
    return [
        t for t in report.transcripts
        if t.best_model is not None
        and t.best_model.locations
        and (min_score is None or (t.best_score or 0) >= min_score)
    ]


def _add_overview_table(document: Document, report: BatchReport, matched_ids: set[int]) -> None:
    heading = document.add_paragraph()
    run = heading.add_run("Overview")
    run.bold = True
    run.font.size = Pt(13)

    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, ("Transcript", "Relevance Score")):
        cell.paragraphs[0].add_run(text).bold = True

    for i, t in enumerate(report.transcripts):
        row = table.add_row()
        name_cell, score_cell = row.cells

        if id(t) in matched_ids:
            _add_internal_link(name_cell.paragraphs[0], t.name, f"t{i}")
        else:
            suffix = "" if t.best_model is not None else " (failed)"
            name_cell.paragraphs[0].add_run(f"{t.name}{suffix}")

        score_run = score_cell.paragraphs[0].add_run(
            str(t.best_score) if t.best_score is not None else "FAILED"
        )
        if t.best_score is None:
            score_run.font.color.rgb = _FAILED_COLOR


def build_docx(report: BatchReport, min_score: int | None = None) -> Document:
    document = Document()
    document.add_heading(report.theme, level=1)
    document.add_paragraph(f"Generated {report.generated_at}")

    matched = _matching_transcripts(report, min_score)
    matched_ids = {id(t) for t in matched}
    _add_overview_table(document, report, matched_ids)

    if not matched:
        document.add_paragraph("No transcripts contained relevant excerpts for this theme.")
        return document

    first = True
    for i, t in enumerate(report.transcripts):
        if id(t) not in matched_ids:
            continue
        if not first:
            _add_horizontal_rule(document)
        first = False

        heading = document.add_heading(f"Transcript: {t.name}", level=2)
        _add_bookmark(heading, f"t{i}", i)
        for location in t.best_model.locations:
            _add_excerpt(document, location)

    return document


def write_docx(report: BatchReport, out_path: str, min_score: int | None = None) -> str:
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    build_docx(report, min_score=min_score).save(out_path)
    return out_path


def write_docx_from_data(
    theme: str,
    transcripts: dict[str, list[tuple[str, dict]]],
    out_path: str,
    min_score: int | None = None,
) -> str:
    """Build and save the Word document directly from in-memory run payloads
    (same shape ``report.write_html_from_data`` takes) — no disk round-trip."""
    report = build_report(theme, transcripts)
    return write_docx(report, out_path, min_score=min_score)


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Regenerate report.docx from the internal analysis cache"
    )
    parser.add_argument("--cache", required=True, help="Path to the .raw_results.json cache file")
    parser.add_argument("--out", default=None, help="Output .docx path (defaults to report.docx next to the cache file)")
    parser.add_argument(
        "--min-score", type=int, default=None,
        help="Only include transcripts whose best score is >= this value (default: include any transcript with at least one located excerpt)",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> None:
    args = parse_args(argv)
    out_path = args.out or os.path.join(os.path.dirname(args.cache) or ".", "report.docx")
    try:
        theme, transcripts = read_cache(args.cache)
    except (FileNotFoundError, OSError) as exc:
        print(str(exc), file=sys.stderr)
        sys.exit(1)
    report = build_report(theme, transcripts)
    write_docx(report, out_path, min_score=args.min_score)
    print(f"[docx] {out_path}")


if __name__ == "__main__":
    main()
